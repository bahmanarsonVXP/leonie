"""
Service de génération de rapports Word pour les dossiers clients.

Ce service crée des rapports automatiques au format .docx montrant:
- Informations client (nom, prénom, type prêt, statut)
- Liste des pièces avec statut (reçu, manquante, non_conforme)
- Statistiques de progression (% complétude)
- Upload automatique sur Google Drive
"""

import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional
from uuid import UUID

import structlog
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.utils.db import (
    get_client_by_id,
    get_pieces_dossier,
    get_courtier_by_id
)
from app.services.drive import DriveManager

logger = structlog.get_logger()


class ReportGenerator:
    """
    Générateur de rapports Word pour suivi des dossiers clients.
    """

    def __init__(self):
        """Initialise le générateur de rapports."""
        self.drive = DriveManager()

    def generate_client_report(
        self,
        client_id: UUID,
        upload_to_drive: bool = True
    ) -> Optional[str]:
        """
        Génère un rapport Word complet pour un client.

        Args:
            client_id: ID du client
            upload_to_drive: Si True, upload le rapport sur Drive

        Returns:
            ID du fichier Drive si upload_to_drive=True, None sinon

        Raises:
            ValueError: Si client non trouvé
        """
        # 1. Récupérer données client
        client = get_client_by_id(client_id)
        if not client:
            raise ValueError(f"Client {client_id} non trouvé")

        courtier = get_courtier_by_id(UUID(client.get('courtier_id')))
        pieces_dossier = get_pieces_dossier(client_id)

        logger.info(
            "Génération rapport client",
            client_id=str(client_id),
            client_nom=f"{client.get('prenom')} {client.get('nom')}",
            nb_pieces=len(pieces_dossier)
        )

        # 2. Créer document Word
        doc = Document()

        # Configuration style global
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # 3. En-tête du rapport
        self._add_header(doc, client, courtier)

        # 4. Statistiques
        stats = self._calculate_statistics(pieces_dossier)
        self._add_statistics(doc, stats)

        # 5. Table des pièces
        self._add_pieces_table(doc, pieces_dossier)

        # 6. Pied de page
        self._add_footer(doc)

        # 7. Sauvegarder temporairement
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            filename = f"RAPPORT_SUIVI_{client.get('nom')}_{client.get('prenom')}.docx"
            file_path = tmp_path / filename

            doc.save(str(file_path))

            logger.info(
                "Rapport Word généré",
                filename=filename,
                file_size=file_path.stat().st_size
            )

            # 8. Upload sur Drive si demandé
            if upload_to_drive:
                client_folder_id = client.get('dossier_drive_id')

                if not client_folder_id:
                    logger.warning(
                        "Client sans dossier Drive, upload impossible",
                        client_id=str(client_id)
                    )
                    return None

                # Nom standard pour le rapport
                drive_filename = "_RAPPORT_SUIVI.docx"

                # Vérifier si rapport existant
                existing_file_id = self.drive.find_file_by_name(
                    drive_filename,
                    client_folder_id
                )

                if existing_file_id:
                    # Mettre à jour le fichier existant
                    logger.info("Mise à jour rapport existant", file_id=existing_file_id)
                    file_id = self.drive.update_file(existing_file_id, file_path)
                else:
                    # Upload nouveau fichier
                    logger.info("Upload nouveau rapport")
                    file_id = self.drive.upload_file(
                        file_path,
                        client_folder_id,
                        drive_filename
                    )

                logger.info(
                    "Rapport uploadé sur Drive",
                    file_id=file_id,
                    folder_id=client_folder_id
                )

                return file_id

        return None

    def _add_header(
        self,
        doc: Document,
        client: Dict,
        courtier: Optional[Dict]
    ) -> None:
        """
        Ajoute l'en-tête du rapport.

        Args:
            doc: Document Word
            client: Données client
            courtier: Données courtier
        """
        # Titre principal
        title = doc.add_heading('Rapport de Suivi - Dossier Client', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Informations client
        doc.add_heading('Informations Client', level=1)

        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Light Grid Accent 1'

        rows_data = [
            ('Nom', client.get('nom', 'N/A')),
            ('Prénom', client.get('prenom', 'N/A')),
            ('Email', client.get('email_principal', 'N/A')),
            ('Type de prêt', client.get('type_pret', 'N/A').replace('_', ' ').title()),
            ('Statut dossier', client.get('statut', 'N/A').replace('_', ' ').title()),
            ('Courtier', f"{courtier.get('prenom', 'N/A')} {courtier.get('nom', 'N/A')}" if courtier else 'N/A')
        ]

        for i, (label, value) in enumerate(rows_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)

            # Style label (gras)
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()  # Espacement

    def _calculate_statistics(self, pieces_dossier: List[Dict]) -> Dict:
        """
        Calcule les statistiques du dossier.

        Args:
            pieces_dossier: Liste des pièces du dossier

        Returns:
            Dict avec statistiques
        """
        total = len(pieces_dossier)

        if total == 0:
            return {
                'total': 0,
                'recues': 0,
                'manquantes': 0,
                'non_conformes': 0,
                'pourcentage_completion': 0
            }

        recues = sum(1 for p in pieces_dossier if p.get('statut') == 'recue')
        manquantes = sum(1 for p in pieces_dossier if p.get('statut') == 'manquante')
        non_conformes = sum(1 for p in pieces_dossier if p.get('statut') == 'non_conforme')

        pourcentage = round((recues / total) * 100, 1) if total > 0 else 0

        return {
            'total': total,
            'recues': recues,
            'manquantes': manquantes,
            'non_conformes': non_conformes,
            'pourcentage_completion': pourcentage
        }

    def _add_statistics(self, doc: Document, stats: Dict) -> None:
        """
        Ajoute les statistiques au rapport.

        Args:
            doc: Document Word
            stats: Statistiques calculées
        """
        doc.add_heading('Statistiques', level=1)

        # Pourcentage de completion (grand)
        completion_para = doc.add_paragraph()
        completion_run = completion_para.add_run(
            f"Taux de complétion : {stats['pourcentage_completion']}%"
        )
        completion_run.font.size = Pt(16)
        completion_run.font.bold = True

        # Couleur selon progression
        if stats['pourcentage_completion'] >= 80:
            completion_run.font.color.rgb = RGBColor(0, 128, 0)  # Vert
        elif stats['pourcentage_completion'] >= 50:
            completion_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            completion_run.font.color.rgb = RGBColor(255, 0, 0)  # Rouge

        completion_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Détails
        doc.add_paragraph(
            f"Total pièces attendues : {stats['total']}\n"
            f"✅ Pièces reçues : {stats['recues']}\n"
            f"⏳ Pièces manquantes : {stats['manquantes']}\n"
            f"❌ Pièces non conformes : {stats['non_conformes']}"
        )

        doc.add_paragraph()  # Espacement

    def _add_pieces_table(
        self,
        doc: Document,
        pieces_dossier: List[Dict]
    ) -> None:
        """
        Ajoute la table des pièces au rapport.

        Args:
            doc: Document Word
            pieces_dossier: Liste des pièces
        """
        doc.add_heading('Détail des Pièces', level=1)

        if not pieces_dossier:
            doc.add_paragraph("Aucune pièce enregistrée pour ce dossier.")
            return

        # Créer table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'

        # En-têtes
        headers = table.rows[0].cells
        headers[0].text = 'Pièce'
        headers[1].text = 'Statut'
        headers[2].text = 'Date réception'
        headers[3].text = 'Commentaire'

        # Style en-têtes (gras)
        for cell in headers:
            cell.paragraphs[0].runs[0].font.bold = True

        # Trier par statut (manquantes en premier)
        statut_ordre = {'manquante': 0, 'non_conforme': 1, 'recue': 2}
        pieces_sorted = sorted(
            pieces_dossier,
            key=lambda p: statut_ordre.get(p.get('statut', 'manquante'), 99)
        )

        # Ajouter lignes
        for piece in pieces_sorted:
            row = table.add_row().cells

            # Nom pièce
            type_piece = piece.get('types_pieces', {})
            nom_piece = type_piece.get('nom', 'Pièce inconnue') if type_piece else 'Pièce inconnue'
            row[0].text = nom_piece

            # Statut
            statut = piece.get('statut', 'manquante')
            statut_text = {
                'recue': '✅ Reçu',
                'manquante': '⏳ Manquante',
                'non_conforme': '❌ Non conforme',
                'non_reconnu': '❓ Non reconnu'
            }.get(statut, statut)
            row[1].text = statut_text

            # Date réception
            date_reception = piece.get('date_reception')
            if date_reception:
                # Formater date ISO en format français
                try:
                    dt = datetime.fromisoformat(date_reception.replace('Z', '+00:00'))
                    row[2].text = dt.strftime('%d/%m/%Y')
                except:
                    row[2].text = 'N/A'
            else:
                row[2].text = 'N/A'

            # Commentaire
            commentaire = piece.get('commentaire', '')
            row[3].text = commentaire if commentaire else '-'

        doc.add_paragraph()  # Espacement

    def _add_footer(self, doc: Document) -> None:
        """
        Ajoute le pied de page au rapport.

        Args:
            doc: Document Word
        """
        doc.add_paragraph()
        doc.add_paragraph('_' * 80)

        footer = doc.add_paragraph()
        footer_run = footer.add_run(
            f"Rapport généré automatiquement par Léonie 🤖\n"
            f"Date de génération : {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
